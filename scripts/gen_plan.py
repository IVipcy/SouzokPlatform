# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x0E, 0x7C, 0x66)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def jp(run, size=None, bold=False, color=None):
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=9.5, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    jp(p.add_run(text), size=size, bold=bold, color=color)
    return p


def cell_lines(cell, lines, size=9.5):
    cell.text = ''
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        prefix = '' if ln.startswith('【') else '・'
        jp(p.add_run(prefix + ln), size=size)


def H(text, level=1):
    h = doc.add_heading(level=level)
    jp(h.add_run(text), color=(BRAND if level <= 2 else ACCENT))
    return h


def P(text, bold=False, size=10.5, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    jp(p.add_run(text), size=size, bold=bold, color=color)
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    jp(p.add_run(text), size=size)
    return p


def make_table(headers, rows, widths=None, header_bg='1F3A5F'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        cell_text(hdr[i], h, bold=True, size=9.5, color=WHITE)
        set_cell_bg(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, list):
                cell_lines(cells[i], val)
            else:
                cell_text(cells[i], val, size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


# ===== タイトル =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
jp(title.add_run('新システム 本番開始（7/3）までの実行プラン'), size=18, bold=True, color=BRAND)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
jp(sub.add_run('作成日 2026/6/22（月）　／　対象期間 6/22(月)〜7/3(金)　／　本番開始 7/3(金)'), size=10, color=GRAY)

# ===== 1 =====
H('1. ゴールと全体方針', 1)
P('7/3(金)の本番開始までに、下記4つを並行で進めて「現場が迷わず使える状態」を作る。', bold=True)
bullet('① トライアル：受注担当・経理担当・事務管理担当の3名が実業務をなぞって検証 → ユーザーFBを反映')
bullet('② データ移行：既存の仕掛中案件を新システムへ移すか判断 → 移すならインポートで効率化')
bullet('③ 研修：全社員向けに2時間の操作研修（録画も残す）')
bullet('④ マニュアル：ロール別マニュアルページを作成し、アプリ内リンクから開けるように')
P('進め方の原則：', bold=True, after=1)
bullet('前半(6/24〜26)でトライアルとFB収集、移行方針の決定を一気に。')
bullet('週末(6/27-28)をFB反映のバッファに充てる。')
bullet('後半(6/29〜7/2)でマニュアル仕上げ・研修・移行実行・最終確認。7/3は伴走サポートに専念。')
bullet('「致命的FB＝本番ブロッカー」だけは必ず本番前に潰す。「あったら良い」は本番後に回してOKと割り切る。')

# ===== 2 =====
H('2. マイルストーン', 1)
make_table(['日付', 'マイルストーン'], [
    ['6/24(水)', 'トライアル開始（3名キックオフ）'],
    ['6/26(金)', 'FB一次締め＋移行する/しないの方針決定'],
    ['6/29(月)', 'ブロッカーFB反映完了'],
    ['6/30(火)', 'マニュアル初版公開（リンク設置）＋インポートのリハーサル'],
    ['7/1(水)', '全社研修（2時間）'],
    ['7/2(木)', '既存案件のインポート本実行＋本番前チェック完了'],
    ['7/3(金)', '本番開始（終日 伴走サポート）'],
], widths=[1.1, 5.2])

# ===== 3 =====
H('3. 日別スケジュール', 1)
make_table(['日付', 'やること', '成果物・ゴール'], [
    ['6/22(月)\n準備', [
        '未適用マイグレーションを本番方針決定の上で適用（前受金行/司・その場で受領補正・契約書類ファイル・返金 ほか）',
        '3名のアカウント・権限・通知設定を確認',
        'トライアル用テスト案件を用意',
        '既存仕掛中案件の棚卸し開始（件数・ステータス・進捗）',
    ], ['本番DB準備', '棚卸し着手']],
    ['6/23(火)\n準備', [
        'ロール別トライアル手順書（各1枚）作成',
        'FB収集フォーム/シート用意（致命/重要/あったら良いで仕分け）',
        'インポート設計：既存データの項目マッピング案を作成',
    ], ['手順書・FBフォーム', 'マッピング案']],
    ['6/24(水)\nトライアルDay1', [
        '3名キックオフ（30分）→各自が実業務をそのまま操作',
        '受注=面談登録→受注内容→請求／事務=対応中→案件状況→タスク着手→調査→受信ボックス／経理=請求発行→入金CSV突合→返金→領収書',
        '終業前に15分FB（フォーム記入）',
    ], ['Day1のFB']],
    ['6/25(木)\nトライアルDay2', [
        '代表案件を各ロール1件、最初から最後まで通す',
        '移行判断会議：基準＝残作業量・事故リスク・二重管理コスト',
        '（目安）完了間近は旧運用で完結／長期の仕掛は新システムへ',
    ], ['移行方針ドラフト']],
    ['6/26(金)\nDay3／FB締め', [
        'FBを「致命/重要/あったら良い」で確定仕分け',
        '移行対象案件の確定リスト＋インポート仕様を確定',
    ], ['FB確定', '移行対象確定']],
    ['6/27(土)〜6/28(日)\nバッファ', [
        'ブロッカーFBの修正実装＆再確認',
        'マニュアルの骨子・キャプチャ取り',
    ], ['ブロッカー解消', 'マニュアル骨子']],
    ['6/29(月)\nFB反映', [
        '重要FBを反映、回帰確認',
        'マニュアル本文執筆（ロール別の主要フロー）',
    ], ['FB反映完了']],
    ['6/30(火)\nマニュアル/移行リハ', [
        'マニュアル初版を公開し、サイドバー/ヘルプにリンク設置',
        'インポートをテスト環境で少数件リハーサル→検証',
        '研修スライド・デモ台本を作成',
    ], ['マニュアル公開', '手順確立']],
    ['7/1(水)\n全社研修', [
        '2時間研修を実施（録画）',
        '研修中の質問をFAQに追記',
    ], ['全社員 操作把握', '録画・FAQ']],
    ['7/2(木)\n移行実行/最終確認', [
        '既存仕掛中案件のインポート本実行→全件検証（やる判断の場合）',
        '本番前チェックリストを消化（migration適用・権限・通知・バックアップ）',
        '予備時間（残課題・追い込み）',
    ], ['移行完了', '本番Go判定']],
    ['7/3(金)\n本番開始', [
        '朝に最終確認、質問窓口を明確化（近くに座って即対応）',
        '立ち上がり伴走、初日トラブルは即時対応',
    ], ['本番稼働']],
], widths=[1.35, 3.65, 1.3])

# ===== 4 =====
H('4. 全社研修（2時間）アジェンダ', 1)
make_table(['時間', '内容'], [
    ['0:00–0:10', '全体像／なぜ新システムか（目的・メリット）'],
    ['0:10–0:40', '受注担当フロー：面談登録→受注内容→請求 のデモ'],
    ['0:40–1:05', '事務管理フロー：対応中→案件状況ボード→タスク着手→調査→受信ボックス→ファイル添付'],
    ['1:05–1:30', '経理フロー：請求発行→入金CSV突合→返金→領収書／通知の流れ'],
    ['1:30–1:45', '共通：案件進捗・通知・マニュアル/ヘルプの使い方'],
    ['1:45–2:00', '質疑応答／本番開始の段取り'],
], widths=[1.1, 5.2])
P('※ 録画して欠席者・後日復習用に共有。研修で出た質問はその場でFAQへ。', size=9.5, color=GRAY)

# ===== 5 =====
H('5. マニュアル方針', 1)
P('まずは「マニュアルページ（ロール別の主要フロー）」をアプリ内リンクで提供。AIチャットは余力次第で次段（FAQが溜まってから）。', bold=True)
P('理由：限られた期間では 2時間研修＋初日伴走＋ロール別ページ の方が確実。AIチャットは構築コストが高く、回答品質の検証も要るため本番後に。', size=10)
P('ページ構成（案）：', bold=True, after=1)
bullet('クイックスタート（ログイン〜基本操作）')
bullet('ロール別フロー：受注担当／事務管理担当／経理担当')
bullet('よくある操作（請求・入金突合・返金・タスク作成・ファイル添付）')
bullet('FAQ（研修・トライアルで出た質問を随時追記）')
bullet('問い合わせ先（社内サポート窓口）')

# ===== 6 =====
H('6. 既存仕掛中案件の移行方針', 1)
P('「全部移す」ありきにしない。案件ごとに費用対効果で判断する。', bold=True)
P('判断基準：', bold=True, after=1)
bullet('残作業が少なく完了間近 → 旧運用のまま完結（移行しない）')
bullet('長期・仕掛が多い → 新システムへインポート（二重管理を避ける）')
bullet('事故リスク（入金・期日管理）が高いもの → 優先的に新システムへ')
P('インポートする場合の最低項目：', bold=True, after=1)
bullet('案件番号（旧番号は lp_case_number 等に保持）・案件名・依頼者・受注担当/管理担当')
bullet('ステータス・受注内容・前受金・完了予定日・進捗メモ（直近の状況）')
P('手順：CSV雛形を用意 → 旧データを記入/出力 → インポートスクリプトで取込 → テスト環境でリハ → 本番取込 → 全件目視検証。', size=10)
P('※「今回は移行せず、旧案件は旧運用で並走」も正当な選択肢。リスクと工数次第で割り切る。', size=9.5, color=GRAY)

# ===== 7 =====
H('7. 本番前チェックリスト（7/2までに全消化）', 1)
for x in [
    '未適用マイグレーションを本番DBへ適用済み（前受金行/司・その場で受領補正・契約書類ファイル・返金 等）',
    '3ロールの権限・表示範囲が正しい',
    '通知（入金確定など）が受注担当・管理担当へ届く',
    '請求→入金CSV突合→返金→領収書 が一通り通る',
    '受信ボックス→各調査表/到着物のファイル参照が動く',
    'データのバックアップを取得済み',
    'マニュアルのリンクが全員から見える',
    '初日サポート体制（窓口・連絡手段）を周知済み',
]:
    bullet(x)

# ===== 8 =====
H('8. リスクとバッファ', 1)
make_table(['リスク', '対策'], [
    ['ブロッカーFBが想定より多い', '6/27-28をバッファに確保。「あったら良い」は本番後送りで割り切る'],
    ['データ移行で不整合', 'テスト環境でリハ→少数件→本番。全件目視検証。最悪は移行せず並走'],
    ['研修で消化不良', '録画＋ロール別ページ＋初日伴走でカバー'],
    ['属人化（あなたに集中）', '3名のトライアル担当を各ロールの一次窓口に育てる'],
], widths=[2.1, 4.2])

# ===== 9 =====
H('9. 役割分担', 1)
make_table(['担当', '役割'], [
    ['あなた（開発）', '修正実装・データ移行・マニュアル作成・研修講師・初日サポート'],
    ['受注担当', '受注フローのトライアル＆FB／研修の受注パート補助'],
    ['事務管理担当', '事務フローのトライアル＆FB／研修の事務パート補助'],
    ['経理担当', '請求・入金・返金のトライアル＆FB／研修の経理パート補助'],
], widths=[1.6, 4.7])

P('', after=2)
P('— 致命的なものだけ確実に潰せば、7/3は十分に間に合います。やり切りましょう。', bold=True, color=ACCENT)

out = 'docs/新システム本番開始プラン_7-3まで.docx'
doc.save(out)
print('SAVED', out)
